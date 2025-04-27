import streamlit as st
from docx import Document
import random
import pandas as pd
from datetime import datetime, timedelta

# Parameters
teams = [
    "Wargrave Wolves", "Oakfield Eagles", "Wollaston Blue", "Parkfield Youth", 
    "Ware Lions", "Ware FC", "West Wight Youth", "Brading Youth", 
    "Brockworth Albion", "Gurnard Rockets"
]
team_colors = {
    "Wargrave Wolves": "Red & Black",
    "Oakfield Eagles": "Royal Blue",
    "Wollaston Blue": "Blue & White",
    "Parkfield Youth": "Red & White",
    "Ware Lions": "Black & White",
    "Ware FC": "Royal Blue & White",
    "West Wight Youth": "Yellow & Black",
    "Brading Youth": "Red & White",
    "Brockworth Albion": "Black & White",
    "Gurnard Rockets": "Green"
}
num_pitches = 3
match_duration = 30  # minutes
seed = 42

competition_days = 2  # Default number of days
start_times = ["13:00", "09:00"]  # Custom start times for Day 1 and Day 2
end_times = ["17:00", "17:00"]   # Custom end times for Day 1 and Day 2

two_legs = True  # Set to True if each team should play twice against each other (home and away)
has_final = True  # Set to True if a final match should be included

def generate_fixture(teams, seed, two_legs=False):
    random.seed(seed)
    random.shuffle(teams)
    fixtures = []
    num_teams = len(teams)

    if num_teams % 2 != 0:  # Add a dummy team if odd number of teams
        teams.append("Bye")
        num_teams += 1

    # Generate first-leg matches
    first_leg = []
    for round_num in range(num_teams - 1):
        round_matches = []
        for i in range(num_teams // 2):
            home = teams[i]
            away = teams[num_teams - 1 - i]
            if home != "Bye" and away != "Bye":
                round_matches.append((home, away))
        first_leg.append(round_matches)
        teams = [teams[0]] + [teams[-1]] + teams[1:-1]  # Rotate the teams

    fixtures.extend(first_leg)

    # Generate second-leg matches (reversing home and away)
    if two_legs:
        second_leg = []
        for round_matches in first_leg:
            reversed_matches = [(away, home) for (home, away) in round_matches]
            second_leg.append(reversed_matches)
        fixtures.extend(second_leg)

    return fixtures


def schedule_matches(fixtures, num_pitches, match_duration, competition_days, start_times=None, end_times=None, has_final=False):
    # Ensure start_times and end_times cover at least days 1 and 2, with default values for additional days
    if start_times is None:
        start_times = ["09:00", "09:00"] + ["09:00"] * (competition_days - 2)
    elif len(start_times) < competition_days:
        start_times += ["09:00"] * (competition_days - len(start_times))

    if end_times is None:
        end_times = ["17:00", "17:00"] + ["17:00"] * (competition_days - 2)
    elif len(end_times) < competition_days:
        end_times += ["17:00"] * (competition_days - len(end_times))

    schedule = []
    active_teams = set()
    matches_per_day = [[] for _ in range(competition_days)]  # To track matches per day
    current_day = 1
    current_time = datetime.strptime(start_times[0], "%H:%M")

    for round_matches in fixtures:
        pitch = 1
        for match in round_matches:
            # Switch days if time exceeds the last match start time for the current day
            if current_time > datetime.strptime(end_times[current_day - 1], "%H:%M"):
                current_day += 1
                if current_day > competition_days:  # Stop if all days are used
                    break
                current_time = datetime.strptime(start_times[current_day - 1], "%H:%M")
                active_teams.clear()

            # Assign matches to available pitches without overlap
            if pitch <= num_pitches and match[0] not in active_teams and match[1] not in active_teams:
                match_entry = {
                    "Date": f"Day {current_day}",
                    "Time": current_time.strftime("%H:%M"),
                    "Pitch": pitch,
                    "Team A": match[0],
                    "Team B": match[1],
                    "Type": "Group",  # Regular match type
                }
                schedule.append(match_entry)
                matches_per_day[current_day - 1].append(match_entry)  # Track match for the day
                active_teams.add(match[0])
                active_teams.add(match[1])
                pitch += 1

            # If all pitches are occupied, move to the next time slot
            if pitch > num_pitches:
                current_time += timedelta(minutes=match_duration)
                pitch = 1
                active_teams.clear()

        # Increment time after processing all matches in the current round
        current_time += timedelta(minutes=match_duration)
        active_teams.clear()

    # Add the final match if `has_final` is True
    if has_final:
        last_match_time = datetime.strptime(schedule[-1]["Time"], "%H:%M")
        final_match_time = last_match_time + timedelta(minutes=30)
        schedule.append({
            "Date": f"Day {current_day}",
            "Time": final_match_time.strftime("%H:%M"),
            "Pitch": 1,  # Default pitch for the final
            "Team A": "Ranked #1 Team",
            "Team B": "Ranked #2 Team",
            "Type": "Final",  # Final match type
        })

    return schedule




def create_word_output(schedule, team_colors, output_file="football_fixture.docx"):
    doc = Document()
    
    doc.add_heading("Football Match Fixture", level=1)
    doc.add_paragraph("ALL MATCHES ARE 11 MINUTES EACH WAY", style='Intense Quote')
    doc.add_heading("Team Colors", level=2)
    
    for team, color in team_colors.items():
        doc.add_paragraph(f"{team} – {color}")
    
    doc.add_heading("Match Schedule", level=2)
    
    for entry in schedule:
        doc.add_paragraph(f"{entry['Date']} | {entry['Time']} | Pitch {entry['Pitch']} | {entry['Team A']} vs {entry['Team B']}")
    
    doc.save(output_file)
    print(f"Word document saved as {output_file}")


#--------- APP
    
# Streamlit App
st.title("Football Fixture Generator")
st.sidebar.header("Settings")


# Other Parameters
num_pitches = st.sidebar.number_input("Number of Pitches", min_value=1, value=3)
match_duration = st.sidebar.number_input("Match Duration (minutes)", min_value=1, value=30)
seed = st.sidebar.number_input("Random Seed", value=42)


# Step 1: Select the Number of Teams
num_teams = st.sidebar.number_input("Number of Teams", min_value=2, max_value=20, value=10)

# Step 2: Input Team Names
teams = []
for i in range(num_teams):
    team_name = st.sidebar.text_input(f"Enter Team {i+1} Name", value=f"Team {i+1}", key=f"team_name_{i}")
    teams.append(team_name)
    
    
    
    
# Step 3: Select Number of Competition Days
competition_days = st.sidebar.number_input("Number of Competition Days", min_value=1, max_value=7, value=2)

# Step 4: Input Start and End Times for Each Day
start_times = []
end_times = []
for day in range(competition_days):
    start_time = st.sidebar.text_input(f"Start Time (Day {day+1})", value="09:00", key=f"start_time_{day}")
    end_time = st.sidebar.text_input(f"End Time (Day {day+1})", value="17:00", key=f"end_time_{day}")
    start_times.append(start_time)
    end_times.append(end_time)

# Step 5: Two Legs Option
two_legs = st.sidebar.checkbox("Play Twice Against Each Other (Two Legs)", value=False)

has_final = st.sidebar.checkbox("Add Final", value=False)

if st.button("Generate Fixture"):
    fixtures = generate_fixture(teams, seed, two_legs)
#    schedule = schedule_matches(fixtures, num_pitches, match_duration, competition_days, start_times, end_times)
    schedule = schedule_matches(fixtures, num_pitches, match_duration, competition_days, start_times, end_times, has_final)
    
    
    
        # Calculate theoretical number of matches
    theoretical_matches = len(teams) * (len(teams) - 1) // 2
    if two_legs:
        theoretical_matches *= 2

    # Display warning if matches generated is fewer than expected
    if len(schedule) < theoretical_matches:
        st.warning(f"Warning: Only {len(schedule)} matches were generated, but {theoretical_matches} matches were expected. This might be due to insufficient time or pitches.")
        
        
        
    # Display the schedule
    df = pd.DataFrame(schedule)
    df = df.reset_index()
    df.rename(columns={'index': 'Match #'}, inplace=True)
    df['Match #'] += 1
    df.set_index('Match #', inplace=True)

    st.dataframe(df)
    
    # Options for Downloads
    st.download_button(
        label="Download Excel",
        data=df.to_csv(index=False),
        file_name="football_fixture.csv",
        mime="text/csv"
    )
    
    create_word_output(schedule, {team: "Color TBD" for team in teams}, output_file="football_fixture.docx")
    with open("football_fixture.docx", "rb") as doc:
        st.download_button(
            label="Download Word Document",
            data=doc,
            file_name="football_fixture.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
